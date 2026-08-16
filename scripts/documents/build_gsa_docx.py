#!/usr/bin/env python3
"""Build the two Kernscale GSA Word working drafts from canonical Markdown."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
GSA_DIR = ROOT / "docs" / "business" / "kernscale" / "gsa"
SOURCE_DIR = GSA_DIR / "application"
OUTPUT_DIR = GSA_DIR / "output"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "666666"
LIGHT_FILL = "F4F6F9"
BORDER = "C9D1D9"
INK = "1F2328"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), BORDER)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def proportional_widths(rows):
    count = len(rows[0])
    if count == 2:
        return [2700, 6660]
    max_lengths = []
    for column in range(count):
        max_lengths.append(max(8, min(42, max(len(row[column]) for row in rows))))
    total = sum(max_lengths)
    raw = [round(TABLE_WIDTH * length / total) for length in max_lengths]
    minimum = 1250 if count <= 4 else 900
    widths = [max(minimum, width) for width in raw]
    while sum(widths) > TABLE_WIDTH:
        index = max(range(count), key=lambda i: widths[i])
        widths[index] -= 10
    while sum(widths) < TABLE_WIDTH:
        index = max(range(count), key=lambda i: max_lengths[i])
        widths[index] += 10
    return widths


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Seite ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def set_run_font(run, *, size=11, bold=None, italic=None, color=INK, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_document(doc, running_title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run(running_title), size=9, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_numbering(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    if kind == "bullet":
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^\)]+\))")


def add_inline(paragraph, text, *, size=11, color=INK):
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            set_run_font(paragraph.add_run(text[cursor:match.start()]), size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), size=size, bold=True, color=color)
        elif token.startswith("`"):
            set_run_font(paragraph.add_run(token[1:-1]), size=size - 0.5, color=color, name="Consolas")
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^\)]+)\)", token).groups()
            run = paragraph.add_run(f"{label} ({url})")
            set_run_font(run, size=size, color=BLUE)
            run.underline = True
        cursor = match.end()
    if cursor < len(text):
        set_run_font(paragraph.add_run(text[cursor:]), size=size, color=color)


def add_markdown_table(doc, rows):
    if len(rows) < 2:
        return
    header = rows[0]
    body = rows[2:] if all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in rows[1]) else rows[1:]
    content = [header] + body
    widths = proportional_widths(content)
    table = doc.add_table(rows=len(content), cols=len(header))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for row_index, source_row in enumerate(content):
        for column_index, value in enumerate(source_row):
            cell = table.cell(row_index, column_index)
            if row_index == 0:
                shade = OxmlElement("w:shd")
                shade.set(qn("w:fill"), LIGHT_FILL)
                cell._tc.get_or_add_tcPr().append(shade)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            add_inline(paragraph, value.strip(), size=9.3, color=INK)
            for run in paragraph.runs:
                run.bold = row_index == 0
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)


def add_cover(doc, title, subtitle, document_type, status):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(28)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("KERNSCALE UG (HAFTUNGSBESCHRÄNKT) I. G."), size=11, bold=True, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run(title), size=25, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    set_run_font(p.add_run(subtitle), size=14, color=DARK_BLUE)

    rows = [
        ["Dokument", document_type],
        ["Antragsteller", "Jonas Knüppel"],
        ["Unternehmenssitz", "Bahnhofstraße 17a, 17094 Burg Stargard"],
        ["Stand", "16.08.2026"],
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [2500, 6860])
    set_table_borders(table)
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, value, size=10)
            if column_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                shade = OxmlElement("w:shd")
                shade.set(qn("w:fill"), LIGHT_FILL)
                cell._tc.get_or_add_tcPr().append(shade)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(18)
    note.paragraph_format.space_after = Pt(8)
    set_run_font(note.add_run(status), size=10, italic=True, color=MUTED)


def parse_markdown(doc, path, *, skip_preamble=True):
    lines = path.read_text(encoding="utf-8").splitlines()
    if skip_preamble:
        start = next((index for index, line in enumerate(lines) if line.startswith("## ")), 0)
        lines = lines[start:]

    index = 0
    in_code = False
    code_lines = []
    active_list_type = None
    active_num_id = None
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.left_indent = Inches(0.18)
                paragraph.paragraph_format.right_indent = Inches(0.18)
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(8)
                paragraph.paragraph_format.line_spacing = 1.0
                p_pr = paragraph._p.get_or_add_pPr()
                shade = OxmlElement("w:shd")
                shade.set(qn("w:fill"), LIGHT_FILL)
                p_pr.append(shade)
                set_run_font(paragraph.add_run("\n".join(code_lines)), size=8.5, color=INK, name="Consolas")
                code_lines = []
                in_code = False
            else:
                in_code = True
            active_list_type = None
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped:
            active_list_type = None
            active_num_id = None
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and lines[index + 1].strip().startswith("|"):
            rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            add_markdown_table(doc, rows)
            active_list_type = None
            active_num_id = None
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)) - 1, 3)
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            add_inline(paragraph, heading.group(2), size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK_BLUE)
            for run in paragraph.runs:
                run.bold = True
            active_list_type = None
            active_num_id = None
            index += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or numbered:
            kind = "bullet" if bullet else "decimal"
            if active_list_type != kind or active_num_id is None:
                active_num_id = add_numbering(doc, kind)
                active_list_type = kind
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.208
            apply_numbering(paragraph, active_num_id)
            add_inline(paragraph, (bullet or numbered).group(1))
            index += 1
            continue
        if stripped.startswith("> "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.15)
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.line_spacing = 1.2
            p_pr = paragraph._p.get_or_add_pPr()
            shade = OxmlElement("w:shd")
            shade.set(qn("w:fill"), LIGHT_FILL)
            p_pr.append(shade)
            add_inline(paragraph, stripped[2:], color=DARK_BLUE)
            active_list_type = None
            active_num_id = None
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if (
                not candidate
                or candidate.startswith("#")
                or candidate.startswith("```")
                or candidate.startswith("|")
                or candidate.startswith("- ")
                or re.match(r"^\d+\.\s+", candidate)
                or candidate.startswith("> ")
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        paragraph = doc.add_paragraph()
        add_inline(paragraph, " ".join(paragraph_lines))
        active_list_type = None
        active_num_id = None


def build_project_idea():
    doc = Document()
    configure_document(doc, "Kernscale | GSA-Projektidee")
    add_cover(
        doc,
        "Kernscale Media Intelligence Loop",
        "Betreute Softwareplattform für kanalübergreifende Medienanalyse und bestätigtes Lernen",
        "Projektidee für das GSA-Gründungsstipendium Mecklenburg-Vorpommern",
        "Arbeitsfassung – Entwicklungszeiträume und einzelne Finanzangaben werden noch ergänzt.",
    )
    parse_markdown(doc, SOURCE_DIR / "02-project-idea.md")
    output = OUTPUT_DIR / "Kernscale-GSA-Projektidee-Arbeitsfassung.docx"
    doc.save(output)
    return output


def build_business_concept():
    doc = Document()
    configure_document(doc, "Kernscale | Unternehmens- und Entwicklungskonzept")
    add_cover(
        doc,
        "Unternehmens- und Entwicklungskonzept",
        "Kernscale Media Intelligence Loop",
        "Anlage zur Projektidee – Produkt, Markt, Umsetzung und Finanzierung",
        "Arbeitsfassung – bestätigte Angaben und offene Entscheidungen sind ausdrücklich getrennt.",
    )
    doc.add_page_break()
    parse_markdown(doc, SOURCE_DIR / "03-business-concept.md")
    doc.add_section(WD_SECTION.NEW_PAGE)
    heading = doc.add_paragraph(style="Heading 1")
    add_inline(heading, "Anlage: Finanzierungs- und Finanzmodell", size=16, color=BLUE)
    for run in heading.runs:
        run.bold = True
    parse_markdown(doc, SOURCE_DIR / "07-financial-model.md")
    output = OUTPUT_DIR / "Kernscale-GSA-Unternehmenskonzept-Arbeitsfassung.docx"
    doc.save(output)
    return output


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    outputs = [build_project_idea(), build_business_concept()]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
